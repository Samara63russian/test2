from __future__ import annotations

import io
import json
import sqlite3
from contextlib import asynccontextmanager
from datetime import date
from pathlib import Path
from typing import Annotated, Any, Literal
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import Depends, FastAPI, Header, HTTPException, Query, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field, field_validator

from reporting_app.database import (
    create_session,
    db,
    hash_password,
    init_db,
    row_to_dict,
    utc_now,
    verify_password,
)


STATIC_DIR = Path(__file__).resolve().parent / "static"


@asynccontextmanager
async def lifespan(_: FastAPI):
    init_db()
    yield


app = FastAPI(
    title="Сводка — справки учреждений",
    version="1.0.0",
    lifespan=lifespan,
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


Role = Literal["admin", "operator", "viewer"]
AnswerType = Literal["text", "textarea", "number", "yes_no", "select", "date"]


class LoginRequest(BaseModel):
    username: str = Field(min_length=2, max_length=80)
    password: str = Field(min_length=6, max_length=256)


class InstitutionInput(BaseModel):
    name: str = Field(min_length=2, max_length=250)
    short_name: str = Field(default="", max_length=100)
    address: str = Field(default="", max_length=500)
    contact_name: str = Field(default="", max_length=200)
    phone: str = Field(default="", max_length=80)
    email: str = Field(default="", max_length=200)


class UserInput(BaseModel):
    username: str = Field(min_length=2, max_length=80)
    password: str = Field(min_length=6, max_length=256)
    full_name: str = Field(min_length=2, max_length=200)
    role: Role = "operator"
    institution_id: int | None = None


class UserUpdate(BaseModel):
    username: str | None = Field(default=None, min_length=2, max_length=80)
    password: str | None = Field(default=None, min_length=6, max_length=256)
    full_name: str | None = Field(default=None, min_length=2, max_length=200)
    role: Role | None = None
    institution_id: int | None = None
    active: bool | None = None


class QuestionInput(BaseModel):
    text: str = Field(min_length=2, max_length=1000)
    help_text: str = Field(default="", max_length=1000)
    answer_type: AnswerType = "text"
    options: list[str] = Field(default_factory=list)
    required: bool = False
    sort_order: int = Field(default=0, ge=0, le=100_000)

    @field_validator("options")
    @classmethod
    def clean_options(cls, values: list[str]) -> list[str]:
        return [value.strip() for value in values if value.strip()]


class ReportInput(BaseModel):
    institution_id: int
    report_date: date
    status: Literal["draft", "submitted"] = "submitted"
    notes: str = Field(default="", max_length=5000)
    answers: dict[str, Any]
    client_uid: str | None = Field(default=None, max_length=100)


def question_dict(row: sqlite3.Row) -> dict[str, Any]:
    item = dict(row)
    item["options"] = json.loads(item.get("options") or "[]")
    item["required"] = bool(item["required"])
    item["active"] = bool(item["active"])
    return item


def institution_dict(row: sqlite3.Row) -> dict[str, Any]:
    item = dict(row)
    item["active"] = bool(item["active"])
    return item


def user_dict(row: sqlite3.Row) -> dict[str, Any]:
    item = dict(row)
    item.pop("password_hash", None)
    item["active"] = bool(item["active"])
    return item


def current_user(
    authorization: Annotated[str | None, Header()] = None,
) -> dict[str, Any]:
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Требуется авторизация")
    token = authorization.split(" ", 1)[1].strip()
    with db() as connection:
        row = connection.execute(
            """
            SELECT u.*
            FROM sessions s
            JOIN users u ON u.id = s.user_id
            WHERE s.token = ? AND s.expires_at > ? AND u.active = 1
            """,
            (token, utc_now()),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=401, detail="Сеанс истёк")
    return user_dict(row)


def require_admin(user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    if user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Доступно только администратору")
    return user


def ensure_institution_access(user: dict[str, Any], institution_id: int) -> None:
    if (
        user["role"] != "admin"
        and user.get("institution_id")
        and user["institution_id"] != institution_id
    ):
        raise HTTPException(status_code=403, detail="Нет доступа к учреждению")


def report_dict(connection: sqlite3.Connection, report_id: int) -> dict[str, Any]:
    report = connection.execute(
        """
        SELECT r.*, i.name AS institution_name, u.full_name AS author_name
        FROM reports r
        JOIN institutions i ON i.id = r.institution_id
        JOIN users u ON u.id = r.created_by
        WHERE r.id = ?
        """,
        (report_id,),
    ).fetchone()
    if not report:
        raise HTTPException(status_code=404, detail="Справка не найдена")
    result = dict(report)
    answers = connection.execute(
        """
        SELECT a.question_id, a.value, q.text AS question_text,
               q.answer_type, q.sort_order
        FROM answers a
        JOIN questions q ON q.id = a.question_id
        WHERE a.report_id = ?
        ORDER BY q.sort_order, q.id
        """,
        (report_id,),
    ).fetchall()
    result["answers"] = [dict(answer) for answer in answers]
    return result


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/api/auth/login")
def login(payload: LoginRequest) -> dict[str, Any]:
    with db() as connection:
        user = connection.execute(
            "SELECT * FROM users WHERE username = ? COLLATE NOCASE AND active = 1",
            (payload.username.strip(),),
        ).fetchone()
        if not user or not verify_password(payload.password, user["password_hash"]):
            raise HTTPException(status_code=401, detail="Неверный логин или пароль")
        token = create_session(connection, user["id"])
        return {"token": token, "user": user_dict(user)}


@app.post("/api/auth/logout", status_code=204)
def logout(
    authorization: Annotated[str | None, Header()] = None,
    _: dict[str, Any] = Depends(current_user),
) -> None:
    token = authorization.split(" ", 1)[1].strip() if authorization else ""
    with db() as connection:
        connection.execute("DELETE FROM sessions WHERE token = ?", (token,))


@app.get("/api/auth/me")
def me(user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    return user


@app.get("/api/bootstrap")
def bootstrap(user: dict[str, Any] = Depends(current_user)) -> dict[str, Any]:
    with db() as connection:
        institutions = connection.execute(
            "SELECT * FROM institutions WHERE active = 1 ORDER BY name"
        ).fetchall()
        questions = connection.execute(
            "SELECT * FROM questions WHERE active = 1 ORDER BY sort_order, id"
        ).fetchall()
    return {
        "user": user,
        "institutions": [institution_dict(row) for row in institutions],
        "questions": [question_dict(row) for row in questions],
    }


@app.get("/api/institutions")
def list_institutions(
    include_inactive: bool = False,
    _: dict[str, Any] = Depends(current_user),
) -> list[dict[str, Any]]:
    query = "SELECT * FROM institutions"
    if not include_inactive:
        query += " WHERE active = 1"
    query += " ORDER BY active DESC, name"
    with db() as connection:
        rows = connection.execute(query).fetchall()
    return [institution_dict(row) for row in rows]


@app.post("/api/institutions", status_code=status.HTTP_201_CREATED)
def create_institution(
    payload: InstitutionInput,
    _: dict[str, Any] = Depends(require_admin),
) -> dict[str, Any]:
    with db() as connection:
        cursor = connection.execute(
            """
            INSERT INTO institutions
                (name, short_name, address, contact_name, phone, email, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.name.strip(),
                payload.short_name.strip(),
                payload.address.strip(),
                payload.contact_name.strip(),
                payload.phone.strip(),
                payload.email.strip(),
                utc_now(),
            ),
        )
        row = connection.execute(
            "SELECT * FROM institutions WHERE id = ?", (cursor.lastrowid,)
        ).fetchone()
    return institution_dict(row)


@app.put("/api/institutions/{institution_id}")
def update_institution(
    institution_id: int,
    payload: InstitutionInput,
    _: dict[str, Any] = Depends(require_admin),
) -> dict[str, Any]:
    with db() as connection:
        cursor = connection.execute(
            """
            UPDATE institutions
            SET name = ?, short_name = ?, address = ?, contact_name = ?,
                phone = ?, email = ?
            WHERE id = ?
            """,
            (
                payload.name.strip(),
                payload.short_name.strip(),
                payload.address.strip(),
                payload.contact_name.strip(),
                payload.phone.strip(),
                payload.email.strip(),
                institution_id,
            ),
        )
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Учреждение не найдено")
        row = connection.execute(
            "SELECT * FROM institutions WHERE id = ?", (institution_id,)
        ).fetchone()
    return institution_dict(row)


@app.delete("/api/institutions/{institution_id}", status_code=204)
def delete_institution(
    institution_id: int,
    _: dict[str, Any] = Depends(require_admin),
) -> None:
    with db() as connection:
        cursor = connection.execute(
            "UPDATE institutions SET active = 0 WHERE id = ?", (institution_id,)
        )
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Учреждение не найдено")


@app.get("/api/questions")
def list_questions(
    include_inactive: bool = False,
    _: dict[str, Any] = Depends(current_user),
) -> list[dict[str, Any]]:
    query = "SELECT * FROM questions"
    if not include_inactive:
        query += " WHERE active = 1"
    query += " ORDER BY active DESC, sort_order, id"
    with db() as connection:
        rows = connection.execute(query).fetchall()
    return [question_dict(row) for row in rows]


@app.post("/api/questions", status_code=status.HTTP_201_CREATED)
def create_question(
    payload: QuestionInput,
    _: dict[str, Any] = Depends(require_admin),
) -> dict[str, Any]:
    if payload.answer_type == "select" and not payload.options:
        raise HTTPException(status_code=422, detail="Добавьте варианты ответа")
    with db() as connection:
        cursor = connection.execute(
            """
            INSERT INTO questions
                (text, help_text, answer_type, options, required, sort_order, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.text.strip(),
                payload.help_text.strip(),
                payload.answer_type,
                json.dumps(payload.options, ensure_ascii=False),
                int(payload.required),
                payload.sort_order,
                utc_now(),
            ),
        )
        row = connection.execute(
            "SELECT * FROM questions WHERE id = ?", (cursor.lastrowid,)
        ).fetchone()
    return question_dict(row)


@app.put("/api/questions/{question_id}")
def update_question(
    question_id: int,
    payload: QuestionInput,
    _: dict[str, Any] = Depends(require_admin),
) -> dict[str, Any]:
    if payload.answer_type == "select" and not payload.options:
        raise HTTPException(status_code=422, detail="Добавьте варианты ответа")
    with db() as connection:
        cursor = connection.execute(
            """
            UPDATE questions
            SET text = ?, help_text = ?, answer_type = ?, options = ?,
                required = ?, sort_order = ?, active = 1
            WHERE id = ?
            """,
            (
                payload.text.strip(),
                payload.help_text.strip(),
                payload.answer_type,
                json.dumps(payload.options, ensure_ascii=False),
                int(payload.required),
                payload.sort_order,
                question_id,
            ),
        )
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Вопрос не найден")
        row = connection.execute(
            "SELECT * FROM questions WHERE id = ?", (question_id,)
        ).fetchone()
    return question_dict(row)


@app.delete("/api/questions/{question_id}", status_code=204)
def delete_question(
    question_id: int,
    _: dict[str, Any] = Depends(require_admin),
) -> None:
    with db() as connection:
        cursor = connection.execute(
            "UPDATE questions SET active = 0 WHERE id = ?", (question_id,)
        )
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="Вопрос не найден")


@app.get("/api/users")
def list_users(
    _: dict[str, Any] = Depends(require_admin),
) -> list[dict[str, Any]]:
    with db() as connection:
        rows = connection.execute(
            """
            SELECT u.*, i.name AS institution_name
            FROM users u
            LEFT JOIN institutions i ON i.id = u.institution_id
            ORDER BY u.active DESC, u.full_name
            """
        ).fetchall()
    return [user_dict(row) for row in rows]


@app.post("/api/users", status_code=status.HTTP_201_CREATED)
def create_user(
    payload: UserInput,
    _: dict[str, Any] = Depends(require_admin),
) -> dict[str, Any]:
    try:
        with db() as connection:
            cursor = connection.execute(
                """
                INSERT INTO users
                    (username, password_hash, full_name, role, institution_id,
                     active, created_at)
                VALUES (?, ?, ?, ?, ?, 1, ?)
                """,
                (
                    payload.username.strip(),
                    hash_password(payload.password),
                    payload.full_name.strip(),
                    payload.role,
                    payload.institution_id,
                    utc_now(),
                ),
            )
            row = connection.execute(
                "SELECT * FROM users WHERE id = ?", (cursor.lastrowid,)
            ).fetchone()
    except sqlite3.IntegrityError as error:
        raise HTTPException(status_code=409, detail="Логин уже используется") from error
    return user_dict(row)


@app.patch("/api/users/{user_id}")
def update_user(
    user_id: int,
    payload: UserUpdate,
    admin: dict[str, Any] = Depends(require_admin),
) -> dict[str, Any]:
    changes = payload.model_dump(exclude_unset=True)
    if not changes:
        raise HTTPException(status_code=422, detail="Нет изменений")
    if user_id == admin["id"] and changes.get("active") is False:
        raise HTTPException(status_code=409, detail="Нельзя отключить свою учётную запись")
    fields: list[str] = []
    values: list[Any] = []
    for key in ("username", "full_name", "role", "institution_id", "active"):
        if key in changes:
            fields.append(f"{key} = ?")
            value = changes[key]
            values.append(int(value) if key == "active" else value)
    if changes.get("password"):
        fields.append("password_hash = ?")
        values.append(hash_password(changes["password"]))
    values.append(user_id)
    try:
        with db() as connection:
            cursor = connection.execute(
                f"UPDATE users SET {', '.join(fields)} WHERE id = ?", values
            )
            if cursor.rowcount == 0:
                raise HTTPException(status_code=404, detail="Пользователь не найден")
            row = connection.execute(
                "SELECT * FROM users WHERE id = ?", (user_id,)
            ).fetchone()
    except sqlite3.IntegrityError as error:
        raise HTTPException(status_code=409, detail="Логин уже используется") from error
    return user_dict(row)


def validate_answers(
    connection: sqlite3.Connection,
    payload: ReportInput,
) -> list[tuple[int, str]]:
    questions = connection.execute(
        "SELECT * FROM questions WHERE active = 1 ORDER BY sort_order, id"
    ).fetchall()
    answers: list[tuple[int, str]] = []
    for question in questions:
        raw_value = payload.answers.get(str(question["id"]), "")
        value = "" if raw_value is None else str(raw_value).strip()
        if payload.status == "submitted" and question["required"] and not value:
            raise HTTPException(
                status_code=422,
                detail=f"Ответьте на обязательный вопрос: {question['text']}",
            )
        if question["answer_type"] == "number" and value:
            try:
                float(value.replace(",", "."))
            except ValueError as error:
                raise HTTPException(
                    status_code=422,
                    detail=f"Ожидается число: {question['text']}",
                ) from error
        if question["answer_type"] == "select" and value:
            options = json.loads(question["options"])
            if value not in options:
                raise HTTPException(status_code=422, detail="Недопустимый вариант ответа")
        answers.append((question["id"], value))
    return answers


@app.post("/api/reports", status_code=status.HTTP_201_CREATED)
def create_report(
    payload: ReportInput,
    user: dict[str, Any] = Depends(current_user),
) -> dict[str, Any]:
    if user["role"] == "viewer":
        raise HTTPException(status_code=403, detail="Роль наблюдателя доступна только для просмотра")
    ensure_institution_access(user, payload.institution_id)
    with db() as connection:
        institution = connection.execute(
            "SELECT id FROM institutions WHERE id = ? AND active = 1",
            (payload.institution_id,),
        ).fetchone()
        if not institution:
            raise HTTPException(status_code=404, detail="Учреждение не найдено")
        if payload.client_uid:
            existing = connection.execute(
                "SELECT id FROM reports WHERE client_uid = ?", (payload.client_uid,)
            ).fetchone()
            if existing:
                return report_dict(connection, existing["id"])
        answers = validate_answers(connection, payload)
        now = utc_now()
        cursor = connection.execute(
            """
            INSERT INTO reports
                (client_uid, institution_id, report_date, status, notes,
                 created_by, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.client_uid,
                payload.institution_id,
                payload.report_date.isoformat(),
                payload.status,
                payload.notes.strip(),
                user["id"],
                now,
                now,
            ),
        )
        connection.executemany(
            "INSERT INTO answers (report_id, question_id, value) VALUES (?, ?, ?)",
            [(cursor.lastrowid, question_id, value) for question_id, value in answers],
        )
        return report_dict(connection, cursor.lastrowid)


@app.put("/api/reports/{report_id}")
def update_report(
    report_id: int,
    payload: ReportInput,
    user: dict[str, Any] = Depends(current_user),
) -> dict[str, Any]:
    if user["role"] == "viewer":
        raise HTTPException(status_code=403, detail="Роль наблюдателя доступна только для просмотра")
    ensure_institution_access(user, payload.institution_id)
    with db() as connection:
        current = connection.execute(
            "SELECT * FROM reports WHERE id = ?", (report_id,)
        ).fetchone()
        if not current:
            raise HTTPException(status_code=404, detail="Справка не найдена")
        ensure_institution_access(user, current["institution_id"])
        answers = validate_answers(connection, payload)
        connection.execute(
            """
            UPDATE reports
            SET institution_id = ?, report_date = ?, status = ?, notes = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                payload.institution_id,
                payload.report_date.isoformat(),
                payload.status,
                payload.notes.strip(),
                utc_now(),
                report_id,
            ),
        )
        connection.execute("DELETE FROM answers WHERE report_id = ?", (report_id,))
        connection.executemany(
            "INSERT INTO answers (report_id, question_id, value) VALUES (?, ?, ?)",
            [(report_id, question_id, value) for question_id, value in answers],
        )
        return report_dict(connection, report_id)


@app.get("/api/reports")
def list_reports(
    institution_id: int | None = None,
    date_from: date | None = None,
    date_to: date | None = None,
    limit: int = Query(default=100, ge=1, le=500),
    user: dict[str, Any] = Depends(current_user),
) -> list[dict[str, Any]]:
    conditions = ["1 = 1"]
    values: list[Any] = []
    effective_institution = institution_id
    if user["role"] != "admin" and user.get("institution_id"):
        effective_institution = user["institution_id"]
    if effective_institution:
        conditions.append("r.institution_id = ?")
        values.append(effective_institution)
    if date_from:
        conditions.append("r.report_date >= ?")
        values.append(date_from.isoformat())
    if date_to:
        conditions.append("r.report_date <= ?")
        values.append(date_to.isoformat())
    values.append(limit)
    with db() as connection:
        rows = connection.execute(
            f"""
            SELECT r.id, r.institution_id, r.report_date, r.status, r.notes,
                   r.created_at, r.updated_at, i.name AS institution_name,
                   u.full_name AS author_name,
                   (SELECT COUNT(*) FROM answers a
                    WHERE a.report_id = r.id AND a.value != '') AS answered_count
            FROM reports r
            JOIN institutions i ON i.id = r.institution_id
            JOIN users u ON u.id = r.created_by
            WHERE {' AND '.join(conditions)}
            ORDER BY r.report_date DESC, r.created_at DESC
            LIMIT ?
            """,
            values,
        ).fetchall()
    return [dict(row) for row in rows]


@app.get("/api/reports/{report_id}")
def get_report(
    report_id: int,
    user: dict[str, Any] = Depends(current_user),
) -> dict[str, Any]:
    with db() as connection:
        result = report_dict(connection, report_id)
    ensure_institution_access(user, result["institution_id"])
    return result


@app.get("/api/reports/{report_id}/document")
def download_report(
    report_id: int,
    user: dict[str, Any] = Depends(current_user),
) -> StreamingResponse:
    with db() as connection:
        report = report_dict(connection, report_id)
    ensure_institution_access(user, report["institution_id"])

    document = Document()
    title = document.add_heading("Итоговая справка", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Учреждение: {report['institution_name']}")
    document.add_paragraph(f"Дата справки: {report['report_date']}")
    document.add_paragraph(f"Ответственный: {report['author_name']}")
    document.add_paragraph(
        f"Статус: {'Отправлена' if report['status'] == 'submitted' else 'Черновик'}"
    )
    document.add_heading("Вопросы и ответы", level=1)
    for index, answer in enumerate(report["answers"], start=1):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{index}. {answer['question_text']}").bold = True
        document.add_paragraph(answer["value"] or "—")
    if report["notes"]:
        document.add_heading("Примечание", level=1)
        document.add_paragraph(report["notes"])
    document.add_paragraph(f"Сформировано: {utc_now()[:10]}")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    filename = quote(f"Справка_{report['report_date']}_{report_id}.docx")
    return StreamingResponse(
        output,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


@app.get("/api/analytics")
def analytics(
    institution_id: int | None = None,
    date_from: date | None = None,
    date_to: date | None = None,
    user: dict[str, Any] = Depends(current_user),
) -> dict[str, Any]:
    conditions = ["r.status = 'submitted'"]
    values: list[Any] = []
    effective_institution = institution_id
    if user["role"] != "admin" and user.get("institution_id"):
        effective_institution = user["institution_id"]
    if effective_institution:
        conditions.append("r.institution_id = ?")
        values.append(effective_institution)
    if date_from:
        conditions.append("r.report_date >= ?")
        values.append(date_from.isoformat())
    if date_to:
        conditions.append("r.report_date <= ?")
        values.append(date_to.isoformat())
    where = " AND ".join(conditions)

    with db() as connection:
        summary = connection.execute(
            f"""
            SELECT COUNT(*) AS total_reports,
                   COUNT(DISTINCT r.institution_id) AS reporting_institutions,
                   MIN(r.report_date) AS first_date,
                   MAX(r.report_date) AS last_date
            FROM reports r WHERE {where}
            """,
            values,
        ).fetchone()
        by_institution = connection.execute(
            f"""
            SELECT i.id, i.name, COUNT(r.id) AS reports
            FROM institutions i
            LEFT JOIN reports r ON r.institution_id = i.id AND {where}
            WHERE i.active = 1
            GROUP BY i.id, i.name
            ORDER BY reports DESC, i.name
            """,
            values,
        ).fetchall()
        trend = connection.execute(
            f"""
            SELECT r.report_date AS date, COUNT(*) AS reports
            FROM reports r
            WHERE {where}
            GROUP BY r.report_date
            ORDER BY r.report_date
            """,
            values,
        ).fetchall()
        total_institutions = connection.execute(
            "SELECT COUNT(*) FROM institutions WHERE active = 1"
        ).fetchone()[0]

    summary_result = dict(summary)
    summary_result["total_institutions"] = total_institutions
    if total_institutions:
        summary_result["coverage_percent"] = round(
            summary_result["reporting_institutions"] / total_institutions * 100
        )
    else:
        summary_result["coverage_percent"] = 0
    return {
        "summary": summary_result,
        "by_institution": [dict(row) for row in by_institution],
        "trend": [dict(row) for row in trend],
    }


app.mount("/", StaticFiles(directory=STATIC_DIR, html=True), name="static")
