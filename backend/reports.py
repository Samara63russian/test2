import io
from datetime import datetime
from typing import List, Optional
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

import models

# Try registering a Cyrillic-compatible font for ReportLab
try:
    # Check standard linux font paths
    font_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf"
    ]
    registered = False
    for fp in font_paths:
        import os
        if os.path.exists(fp):
            pdfmetrics.registerFont(TTFont('DejaVuSans', fp))
            pdf_font_name = 'DejaVuSans'
            registered = True
            break
    if not registered:
        pdf_font_name = 'Helvetica'
except Exception as e:
    pdf_font_name = 'Helvetica'

def generate_pdf_report(report: models.InspectionReport, institution: models.Institution, answers: List[models.Answer]) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'RussianTitle',
        parent=styles['Heading1'],
        fontName=pdf_font_name,
        fontSize=15,
        leading=18,
        alignment=1, # Center
        textColor=colors.HexColor('#1e293b'),
        spaceAfter=12
    )

    h2_style = ParagraphStyle(
        'RussianH2',
        parent=styles['Heading2'],
        fontName=pdf_font_name,
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#0f766e'),
        spaceBefore=10,
        spaceAfter=6
    )

    normal_style = ParagraphStyle(
        'RussianNormal',
        parent=styles['Normal'],
        fontName=pdf_font_name,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#334155')
    )

    bold_style = ParagraphStyle(
        'RussianBold',
        parent=normal_style,
        fontName=pdf_font_name,
        fontStyle='Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#0f172a')
    )

    story = []

    # Title
    story.append(Paragraph("СВОДНАЯ ИНФОРМАЦИОННО-АНАЛИТИЧЕСКАЯ СПРАВКА", title_style))
    story.append(Paragraph(f"по результатам обследования учреждения: <b>{institution.name}</b>", normal_style))
    story.append(Spacer(1, 10))

    # General info table
    insp_date_str = report.inspection_date.strftime('%d.%m.%Y %H:%M') if report.inspection_date else "-"
    inspector_str = report.inspector.full_name if report.inspector else "Не указан"
    
    info_data = [
        [Paragraph("<b>Наименование учреждения:</b>", normal_style), Paragraph(institution.name or "-", normal_style)],
        [Paragraph("<b>Категория / Отрасль:</b>", normal_style), Paragraph(institution.category or "-", normal_style)],
        [Paragraph("<b>Адрес:</b>", normal_style), Paragraph(institution.address or "-", normal_style)],
        [Paragraph("<b>Руководитель:</b>", normal_style), Paragraph(institution.head_name or "-", normal_style)],
        [Paragraph("<b>Дата проведения проверки:</b>", normal_style), Paragraph(insp_date_str, normal_style)],
        [Paragraph("<b>Проверяющий / Эксперт:</b>", normal_style), Paragraph(inspector_str, normal_style)],
        [Paragraph("<b>Итоговый балл соответствия:</b>", bold_style), Paragraph(f"<b>{report.score:.1f}%</b>", bold_style)],
        [Paragraph("<b>Статус справки:</b>", normal_style), Paragraph("Завершена и утверждена" if report.status == "completed" else report.status, normal_style)],
    ]

    info_table = Table(info_data, colWidths=[160, 360])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 12))

    # Summary Text & Recommendations
    if report.summary_text:
        story.append(Paragraph("Краткое резюме и заключение:", h2_style))
        story.append(Paragraph(report.summary_text.replace('\n', '<br/>'), normal_style))
        story.append(Spacer(1, 8))

    if report.recommendations:
        story.append(Paragraph("Рекомендации и план устранения замечаний:", h2_style))
        story.append(Paragraph(report.recommendations.replace('\n', '<br/>'), normal_style))
        story.append(Spacer(1, 10))

    # Questions & Answers Table
    story.append(Paragraph("Детализация ответов на вопросы чек-листа:", h2_style))
    
    table_rows = [
        [
            Paragraph("<b>№</b>", bold_style),
            Paragraph("<b>Вопрос / Критерий оценки</b>", bold_style),
            Paragraph("<b>Ответ</b>", bold_style),
            Paragraph("<b>Соотв.</b>", bold_style),
            Paragraph("<b>Комментарий</b>", bold_style)
        ]
    ]

    for idx, ans in enumerate(answers, start=1):
        q_text = ans.question.text if ans.question else f"Вопрос #{ans.question_id}"
        q_code = f"[{ans.question.code}] " if ans.question and ans.question.code else ""
        full_q = f"{q_code}{q_text}"
        val_str = str(ans.value if ans.value is not None else "-")
        if val_str.lower() == "true":
            val_str = "Да"
        elif val_str.lower() == "false":
            val_str = "Нет"
            
        comp_str = "Да" if ans.is_compliant is True else ("Нет" if ans.is_compliant is False else "—")
        comm_str = ans.comment or ""

        table_rows.append([
            Paragraph(str(idx), normal_style),
            Paragraph(full_q, normal_style),
            Paragraph(val_str, normal_style),
            Paragraph(comp_str, normal_style),
            Paragraph(comm_str, normal_style)
        ])

    qa_table = Table(table_rows, colWidths=[24, 230, 80, 46, 140])
    qa_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0f766e')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('ALIGN', (0,0), (0,-1), 'CENTER'),
        ('ALIGN', (3,0), (3,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
    ]))
    story.append(qa_table)
    story.append(Spacer(1, 16))

    # Signatures block
    sig_data = [
        [
            Paragraph("<b>Проверяющий эксперт:</b>", normal_style),
            Paragraph("____________________ / " + inspector_str + " /", normal_style)
        ],
        [
            Paragraph("<b>Представитель учреждения:</b>", normal_style),
            Paragraph("____________________ / " + (institution.head_name or "____________________") + " /", normal_style)
        ]
    ]
    sig_table = Table(sig_data, colWidths=[200, 320])
    sig_table.setStyle(TableStyle([
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(KeepTogether(sig_table))

    doc.build(story)
    return buffer.getvalue()

def generate_docx_report(report: models.InspectionReport, institution: models.Institution, answers: List[models.Answer]) -> bytes:
    doc = docx.Document()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("СВОДНАЯ ИНФОРМАЦИОННО-АНАЛИТИЧЕСКАЯ СПРАВКА\nпо результатам обследования учреждения")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"«{institution.name}»")
    r_sub.bold = True
    r_sub.font.size = Pt(13)

    # General info table
    doc.add_heading("1. Общие сведения об учреждении и проверке", level=2)
    
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    insp_date_str = report.inspection_date.strftime('%d.%m.%Y %H:%M') if report.inspection_date else "-"
    inspector_str = report.inspector.full_name if report.inspector else "Не указан"

    fields = [
        ("Наименование учреждения", institution.name or "-"),
        ("Отрасль / Категория", institution.category or "-"),
        ("Адрес местонахождения", institution.address or "-"),
        ("Руководитель учреждения", institution.head_name or "-"),
        ("Контактный телефон", institution.phone or "-"),
        ("Дата и время обследования", insp_date_str),
        ("Ответственный эксперт", inspector_str),
        ("Итоговый процент соответствия", f"{report.score:.1f}%"),
        ("Статус документа", "Завершена" if report.status == "completed" else report.status)
    ]

    for label, val in fields:
        row = table.add_row()
        c0 = row.cells[0]
        c1 = row.cells[1]
        c0.paragraphs[0].add_run(label).bold = True
        c1.paragraphs[0].add_run(str(val))

    # Summary
    if report.summary_text:
        doc.add_heading("2. Итоговое резюме и выводы", level=2)
        doc.add_paragraph(report.summary_text)

    # Recommendations
    if report.recommendations:
        doc.add_heading("3. Предписания и рекомендации", level=2)
        doc.add_paragraph(report.recommendations)

    # Checklist table
    doc.add_heading("4. Результаты оценки по опросному листу", level=2)
    qa_tbl = doc.add_table(rows=1, cols=5)
    qa_tbl.style = 'Table Grid'
    
    hdr_cells = qa_tbl.rows[0].cells
    hdr_titles = ["№", "Критерий / Вопрос", "Ответ", "Соответствие", "Замечания / Комментарии"]
    for i, t in enumerate(hdr_titles):
        hdr_cells[i].paragraphs[0].add_run(t).bold = True
        
    for idx, ans in enumerate(answers, start=1):
        row = qa_tbl.add_row()
        q_text = ans.question.text if ans.question else f"Вопрос #{ans.question_id}"
        q_code = f"[{ans.question.code}] " if ans.question and ans.question.code else ""
        full_q = f"{q_code}{q_text}"
        val_str = str(ans.value if ans.value is not None else "-")
        if val_str.lower() == "true":
            val_str = "Да"
        elif val_str.lower() == "false":
            val_str = "Нет"
        comp_str = "Да" if ans.is_compliant is True else ("Нет" if ans.is_compliant is False else "—")
        
        row.cells[0].paragraphs[0].add_run(str(idx))
        row.cells[1].paragraphs[0].add_run(full_q)
        row.cells[2].paragraphs[0].add_run(val_str)
        row.cells[3].paragraphs[0].add_run(comp_str)
        row.cells[4].paragraphs[0].add_run(ans.comment or "")

    doc.add_paragraph("\n\n")
    p_sig = doc.add_paragraph()
    p_sig.add_run(f"Проверяющий: ___________________ / {inspector_str} /\n\n")
    p_sig.add_run(f"Руководитель учреждения: ___________________ / {institution.head_name or ''} /")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def generate_excel_report(reports_list: List[models.InspectionReport]) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Сводная ведомость справок"

    headers = [
        "ID", "Дата проверки", "Учреждение", "Категория", "Адрес", "Руководитель",
        "Проверяющий", "Статус", "Итоговый балл (%)", "Заголовок справки", "Резюме"
    ]
    
    ws.append(headers)
    
    # Styling header
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="0F766E", end_color="0F766E", fill_type="solid")
    border_thin = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for r in reports_list:
        inst = r.institution
        date_str = r.inspection_date.strftime("%d.%m.%Y %H:%M") if r.inspection_date else ""
        ws.append([
            r.id,
            date_str,
            inst.name if inst else "-",
            inst.category if inst else "-",
            inst.address if inst else "-",
            inst.head_name if inst else "-",
            r.inspector.full_name if r.inspector else "-",
            r.status,
            round(r.score, 1),
            r.title,
            r.summary_text or ""
        ])

    for row in ws.iter_rows(min_row=2, max_row=len(reports_list) + 1, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.border = border_thin
            cell.alignment = Alignment(vertical="center")

    # Column widths
    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
