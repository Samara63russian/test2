from io import BytesIO

from docx import Document
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.models import Institution, Question, Report, User


def generate_report_docx(db: Session, report: Report) -> BytesIO:
    institution = db.query(Institution).filter(Institution.id == report.institution_id).first()
    author = db.query(User).filter(User.id == report.author_id).first()
    questions = {q.id: q for q in db.query(Question).all()}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_heading("Сводная справка", level=0)
    title.alignment = 1

    doc.add_paragraph(f"Учреждение: {institution.name if institution else '—'}")
    if institution and institution.address:
        doc.add_paragraph(f"Адрес: {institution.address}")
    doc.add_paragraph(f"Дата справки: {report.report_date.strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Автор: {author.full_name or author.username if author else '—'}")
    doc.add_paragraph(f"Статус: {report.status}")
    doc.add_paragraph("")

    answers_by_q = {a.question_id: a.answer_text for a in report.answers}
    current_category = None

    for question in sorted(questions.values(), key=lambda q: (q.category, q.sort_order, q.id)):
        if not question.is_active:
            continue
        if question.category != current_category:
            current_category = question.category
            doc.add_heading(current_category, level=1)
        doc.add_paragraph(f"Вопрос: {question.text}")
        answer = answers_by_q.get(question.id, "")
        doc.add_paragraph(f"Ответ: {answer or '—'}")
        doc.add_paragraph("")

    if report.notes:
        doc.add_heading("Примечания", level=1)
        doc.add_paragraph(report.notes)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
