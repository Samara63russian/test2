from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm
from fpdf import FPDF
from sqlalchemy.orm import Session

from ..models import AppSetting, Question, QuestionCategory, Report

FONT_DIR = Path(__file__).resolve().parent.parent / "fonts"
SYSTEM_FONT_DIR = Path("/usr/share/fonts/truetype/dejavu")


def _settings(db: Session) -> dict[str, str]:
    return {s.key: s.value for s in db.query(AppSetting).all()}


def _grouped(db: Session, report: Report) -> list[tuple[str, list[tuple[str, str]]]]:
    answers = {a.question_id: a.value for a in report.answers}
    categories = (
        db.query(QuestionCategory)
        .filter(QuestionCategory.is_active.is_(True))
        .order_by(QuestionCategory.sort_order, QuestionCategory.id)
        .all()
    )
    grouped = []
    for cat in categories:
        questions = (
            db.query(Question)
            .filter(Question.category_id == cat.id, Question.is_active.is_(True))
            .order_by(Question.sort_order, Question.id)
            .all()
        )
        rows = [(q.text, answers.get(q.id, "") or "—") for q in questions]
        if rows:
            grouped.append((cat.name, rows))
    return grouped


def build_docx(db: Session, report: Report) -> bytes:
    s = _settings(db)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(16)

    header = doc.add_paragraph(s.get("org_header", ""))
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.bold = True
        run.font.size = Pt(11)

    title = doc.add_paragraph(s.get("document_title", "СВОДНАЯ СПРАВКА"))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    city = doc.add_paragraph(s.get("org_city", ""))
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER

    inst = report.institution
    meta = doc.add_paragraph()
    meta.add_run("Учреждение: ").bold = True
    meta.add_run(inst.name if inst else "—")
    meta2 = doc.add_paragraph()
    meta2.add_run("Код: ").bold = True
    meta2.add_run((inst.code if inst else "") or "—")
    meta2.add_run("    Дата справки: ").bold = True
    meta2.add_run(report.report_date.strftime("%d.%m.%Y"))
    meta3 = doc.add_paragraph()
    meta3.add_run("Адрес: ").bold = True
    meta3.add_run((inst.address if inst else "") or "—")
    meta3.add_run("    Руководитель: ").bold = True
    meta3.add_run((inst.head_name if inst else "") or "—")
    meta4 = doc.add_paragraph()
    meta4.add_run("Составил: ").bold = True
    meta4.add_run((report.user.full_name if report.user else "") or "—")
    meta4.add_run("    Статус: ").bold = True
    meta4.add_run("утверждена" if report.status == "submitted" else "черновик")

    for cat_name, rows in _grouped(db, report):
        heading = doc.add_paragraph(cat_name)
        heading.runs[0].bold = True
        heading.runs[0].font.size = Pt(13)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Вопрос"
        hdr[1].text = "Ответ"
        for q, a in rows:
            row = table.add_row().cells
            row[0].text = q
            row[1].text = a

    doc.add_paragraph("")
    doc.add_paragraph(s.get("footer_text", ""))
    sign = doc.add_paragraph("_________________ / _________________")
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cap = doc.add_paragraph("подпись                расшифровка")
    cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _font_file(name: str) -> Path | None:
    for folder in (FONT_DIR, SYSTEM_FONT_DIR):
        path = folder / name
        if path.exists():
            return path
    return None


class ReportPDF(FPDF):
    def __init__(self, footer_text: str):
        super().__init__()
        self.footer_text = footer_text
        regular = _font_file("DejaVuSans.ttf")
        bold = _font_file("DejaVuSans-Bold.ttf") or regular
        if regular:
            self.add_font("DejaVu", "", str(regular))
            self.add_font("DejaVu", "B", str(bold))
            self.font_name = "DejaVu"
        else:
            self.font_name = "Helvetica"

    def footer(self):
        self.set_y(-15)
        self.set_font(self.font_name, "", 8)
        self.cell(0, 8, f"{self.footer_text}  стр. {self.page_no()}", align="C")

    def block(self, text: str, size: float = 11, bold: bool = False, align: str = "L", height: float = 6):
        self.set_font(self.font_name, "B" if bold else "", size)
        self.set_x(self.l_margin)
        self.multi_cell(self.epw, height, text or " ", align=align, new_x="LMARGIN", new_y="NEXT")


def build_pdf(db: Session, report: Report) -> bytes:
    s = _settings(db)
    pdf = ReportPDF(s.get("footer_text", ""))
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    for line in s.get("org_header", "").split("\n"):
        pdf.block(line, size=11, bold=True, align="C")
    pdf.ln(4)
    pdf.block(s.get("document_title", "СВОДНАЯ СПРАВКА"), size=16, bold=True, align="C", height=8)
    pdf.block(s.get("org_city", ""), size=11, align="C")
    pdf.ln(4)

    inst = report.institution
    lines = [
        f"Учреждение: {inst.name if inst else '—'}",
        f"Код: {(inst.code if inst else '') or '—'}    Дата справки: {report.report_date.strftime('%d.%m.%Y')}",
        f"Адрес: {(inst.address if inst else '') or '—'}    Руководитель: {(inst.head_name if inst else '') or '—'}",
        f"Составил: {(report.user.full_name if report.user else '') or '—'}    Статус: {'утверждена' if report.status == 'submitted' else 'черновик'}",
    ]
    for line in lines:
        pdf.block(line, size=11)
    pdf.ln(2)

    for cat_name, rows in _grouped(db, report):
        pdf.block(cat_name, size=13, bold=True, height=8)
        for q, a in rows:
            pdf.block(q, size=10, bold=True, height=5)
            pdf.block(a, size=10, height=5)
            pdf.ln(1)

    pdf.ln(8)
    pdf.block("_________________ / _________________", size=11, align="R")
    pdf.block("подпись                расшифровка", size=10, align="R")
    return bytes(pdf.output())
