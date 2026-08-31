from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from sqlalchemy.orm import Session

from .models import AnswerOption, Question, Report


def _option_labels(db: Session, option_ids: list[int]) -> str:
    if not option_ids:
        return ""
    options = db.query(AnswerOption).filter(AnswerOption.id.in_(option_ids)).all()
    by_id = {o.id: o.text for o in options}
    return ", ".join(by_id[i] for i in option_ids if i in by_id)


def parse_option_ids(raw: str) -> list[int]:
    if not raw:
        return []
    return [int(x) for x in raw.split(",") if x.strip().isdigit()]


def build_docx(db: Session, report: Report) -> bytes:
    doc = Document()
    title = doc.add_heading("Сводная справка", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run("Учреждение: ").bold = True
    meta.add_run(report.institution.name if report.institution else "—")

    meta2 = doc.add_paragraph()
    meta2.add_run("Дата справки: ").bold = True
    meta2.add_run(report.report_date.isoformat())

    meta3 = doc.add_paragraph()
    meta3.add_run("Статус: ").bold = True
    meta3.add_run("Отправлена" if report.status == "submitted" else "Черновик")

    if report.author:
        meta4 = doc.add_paragraph()
        meta4.add_run("Автор: ").bold = True
        meta4.add_run(report.author.full_name or report.author.username)

    if report.notes:
        notes = doc.add_paragraph()
        notes.add_run("Примечание: ").bold = True
        notes.add_run(report.notes)

    doc.add_heading("Ответы на вопросы", level=1)

    answers_by_q = {a.question_id: a for a in report.answers}
    questions = (
        db.query(Question)
        .filter(Question.is_active.is_(True))
        .order_by(Question.sort_order, Question.id)
        .all()
    )

    for idx, question in enumerate(questions, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {question.text}")
        run.bold = True
        run.font.size = Pt(11)

        answer = answers_by_q.get(question.id)
        value = "—"
        if answer:
            option_ids = parse_option_ids(answer.option_ids)
            if option_ids:
                value = _option_labels(db, option_ids)
                if answer.value_text:
                    value = f"{value}; {answer.value_text}"
            else:
                value = answer.value_text or "—"

        a = doc.add_paragraph(value)
        a.paragraph_format.space_after = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_xlsx(db: Session, reports: list[Report]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Справки"

    questions = (
        db.query(Question)
        .filter(Question.is_active.is_(True))
        .order_by(Question.sort_order, Question.id)
        .all()
    )

    headers = ["ID", "Учреждение", "Дата", "Статус", "Автор", "Примечание"] + [
        q.text for q in questions
    ]
    ws.append(headers)

    for report in reports:
        answers_by_q = {a.question_id: a for a in report.answers}
        row = [
            report.id,
            report.institution.name if report.institution else "",
            report.report_date.isoformat(),
            report.status,
            (report.author.full_name or report.author.username) if report.author else "",
            report.notes,
        ]
        for question in questions:
            answer = answers_by_q.get(question.id)
            if not answer:
                row.append("")
                continue
            option_ids = parse_option_ids(answer.option_ids)
            if option_ids:
                value = _option_labels(db, option_ids)
                if answer.value_text:
                    value = f"{value}; {answer.value_text}"
                row.append(value)
            else:
                row.append(answer.value_text)
        ws.append(row)

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
