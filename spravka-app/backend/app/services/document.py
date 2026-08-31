from io import BytesIO
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from openpyxl import Workbook
from sqlalchemy.orm import Session

from ..models import Report, Question


def build_report_docx(db: Session, report: Report) -> BytesIO:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report.title.upper())
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Учреждение: {report.institution.name}\n"
        f"Дата: {report.report_date}\n"
        f"Статус: {report.status}"
    )

    doc.add_paragraph("")

    questions = {
        q.id: q
        for q in db.query(Question).order_by(Question.sort_order, Question.id).all()
    }
    answers_by_q = {a.question_id: a for a in report.answers}

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Вопрос"
    hdr[2].text = "Ответ"

    for idx, qid in enumerate(sorted(questions.keys(), key=lambda i: questions[i].sort_order), start=1):
        q = questions[qid]
        if not q.is_active and qid not in answers_by_q:
            continue
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = q.text
        ans = answers_by_q.get(qid)
        if not ans:
            row[2].text = "—"
        elif ans.answer_option_id and ans.answer_option:
            row[2].text = ans.answer_option.text
        else:
            row[2].text = ans.text_value or "—"

    if report.notes:
        doc.add_paragraph("")
        notes = doc.add_paragraph()
        notes.add_run("Примечания: ").bold = True
        notes.add_run(report.notes)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(
        f"Сформировано: {datetime.utcnow().strftime('%d.%m.%Y %H:%M')} UTC"
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_report_xlsx(db: Session, report: Report) -> BytesIO:
    wb = Workbook()
    ws = wb.active
    ws.title = "Справка"
    ws.append(["Учреждение", report.institution.name])
    ws.append(["Дата", report.report_date])
    ws.append(["Название", report.title])
    ws.append(["Статус", report.status])
    ws.append([])
    ws.append(["№", "Категория", "Вопрос", "Ответ"])

    questions = db.query(Question).order_by(Question.sort_order, Question.id).all()
    answers_by_q = {a.question_id: a for a in report.answers}
    for idx, q in enumerate(questions, start=1):
        ans = answers_by_q.get(q.id)
        if ans and ans.answer_option:
            value = ans.answer_option.text
        elif ans:
            value = ans.text_value
        else:
            value = ""
        ws.append([idx, q.category, q.text, value])

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
